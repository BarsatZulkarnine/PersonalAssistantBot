"""
Document Loaders - Support Multiple File Formats

Loaders for PDF, DOCX, TXT, Markdown, HTML, etc.
"""

import os
from pathlib import Path
from typing import Optional
from modules.rag.base import DocumentLoader, Document, DocumentType
from utils.logger import get_logger

logger = get_logger('rag.loaders')

# ===== TEXT LOADER =====

class TextLoader(DocumentLoader):
    """Load plain text and markdown files"""
    
    def can_load(self, file_path: str) -> bool:
        """Check if file is text or markdown"""
        ext = Path(file_path).suffix.lower()
        return ext in ['.txt', '.md', '.markdown']
    
    def load(self, file_path: str) -> Document:
        """Load text file"""
        path = Path(file_path)
        
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Determine type
        ext = path.suffix.lower()
        doc_type = DocumentType.MARKDOWN if ext in ['.md', '.markdown'] else DocumentType.TXT
        
        doc = Document(
            file_path=str(path.absolute()),
            file_name=path.name,
            file_type=doc_type,
            title=path.stem,
            content=content,
            file_size_bytes=path.stat().st_size,
            metadata={'extension': ext}
        )
        
        logger.info(f"Loaded text: {path.name} ({len(content)} chars)")
        return doc


# ===== PDF LOADER =====

class PDFLoader(DocumentLoader):
    """Load PDF files"""
    
    def __init__(self):
        self.available = self._check_dependencies()
    
    def _check_dependencies(self) -> bool:
        """Check if PDF libraries are available"""
        try:
            import PyPDF2
            return True
        except ImportError:
            logger.warning("PyPDF2 not installed. Install with: pip install PyPDF2")
            return False
    
    def can_load(self, file_path: str) -> bool:
        """Check if file is PDF"""
        return Path(file_path).suffix.lower() == '.pdf' and self.available
    
    def load(self, file_path: str) -> Document:
        """Load PDF file"""
        if not self.available:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        import PyPDF2
        
        path = Path(file_path)
        text_content = []
        
        try:
            with open(path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num}: {e}")
                
                # Get metadata
                metadata = {
                    'num_pages': len(pdf_reader.pages),
                    'extension': '.pdf'
                }
                
                # Try to get PDF metadata
                if pdf_reader.metadata:
                    if '/Title' in pdf_reader.metadata:
                        metadata['pdf_title'] = pdf_reader.metadata['/Title']
                    if '/Author' in pdf_reader.metadata:
                        metadata['pdf_author'] = pdf_reader.metadata['/Author']
        
        except Exception as e:
            logger.error(f"Failed to load PDF {path}: {e}")
            raise
        
        content = "\n\n".join(text_content)
        
        doc = Document(
            file_path=str(path.absolute()),
            file_name=path.name,
            file_type=DocumentType.PDF,
            title=metadata.get('pdf_title', path.stem),
            content=content,
            file_size_bytes=path.stat().st_size,
            metadata=metadata
        )
        
        logger.info(f"Loaded PDF: {path.name} ({len(content)} chars, {metadata.get('num_pages', 0)} pages)")
        return doc


# ===== DOCX LOADER =====

class DOCXLoader(DocumentLoader):
    """Load Microsoft Word documents"""
    
    def __init__(self):
        self.available = self._check_dependencies()
    
    def _check_dependencies(self) -> bool:
        """Check if python-docx is available"""
        try:
            import docx
            return True
        except ImportError:
            logger.warning("python-docx not installed. Install with: pip install python-docx")
            return False
    
    def can_load(self, file_path: str) -> bool:
        """Check if file is DOCX"""
        return Path(file_path).suffix.lower() == '.docx' and self.available
    
    def load(self, file_path: str) -> Document:
        """Load DOCX file"""
        if not self.available:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        import docx
        
        path = Path(file_path)
        
        try:
            doc_obj = docx.Document(path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc_obj.paragraphs if para.text.strip()]
            content = "\n\n".join(paragraphs)
            
            # Get metadata
            metadata = {
                'num_paragraphs': len(doc_obj.paragraphs),
                'extension': '.docx'
            }
            
            # Try to get core properties
            try:
                if doc_obj.core_properties.title:
                    metadata['docx_title'] = doc_obj.core_properties.title
                if doc_obj.core_properties.author:
                    metadata['docx_author'] = doc_obj.core_properties.author
            except:
                pass
        
        except Exception as e:
            logger.error(f"Failed to load DOCX {path}: {e}")
            raise
        
        doc = Document(
            file_path=str(path.absolute()),
            file_name=path.name,
            file_type=DocumentType.DOCX,
            title=metadata.get('docx_title', path.stem),
            content=content,
            file_size_bytes=path.stat().st_size,
            metadata=metadata
        )
        
        logger.info(f"Loaded DOCX: {path.name} ({len(content)} chars, {metadata.get('num_paragraphs', 0)} paragraphs)")
        return doc


# ===== HTML LOADER =====

class HTMLLoader(DocumentLoader):
    """Load HTML files"""
    
    def __init__(self):
        self.available = self._check_dependencies()
    
    def _check_dependencies(self) -> bool:
        """Check if BeautifulSoup is available"""
        try:
            from bs4 import BeautifulSoup
            return True
        except ImportError:
            logger.warning("beautifulsoup4 not installed. Install with: pip install beautifulsoup4")
            return False
    
    def can_load(self, file_path: str) -> bool:
        """Check if file is HTML"""
        ext = Path(file_path).suffix.lower()
        return ext in ['.html', '.htm'] and self.available
    
    def load(self, file_path: str) -> Document:
        """Load HTML file"""
        if not self.available:
            raise ImportError("beautifulsoup4 not installed. Install with: pip install beautifulsoup4")
        
        from bs4 import BeautifulSoup
        
        path = Path(file_path)
        
        with open(path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(['script', 'style']):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        content = '\n'.join(chunk for chunk in chunks if chunk)
        
        # Get title
        title = path.stem
        title_tag = soup.find('title')
        if title_tag:
            title = title_tag.string
        
        doc = Document(
            file_path=str(path.absolute()),
            file_name=path.name,
            file_type=DocumentType.HTML,
            title=title,
            content=content,
            file_size_bytes=path.stat().st_size,
            metadata={'extension': path.suffix}
        )
        
        logger.info(f"Loaded HTML: {path.name} ({len(content)} chars)")
        return doc


# ===== LOADER REGISTRY =====

class LoaderRegistry:
    """Registry for document loaders"""
    
    def __init__(self):
        self.loaders = [
            TextLoader(),
            PDFLoader(),
            DOCXLoader(),
            HTMLLoader()
        ]
        
        logger.info(f"LoaderRegistry initialized with {len(self.loaders)} loaders")
    
    def get_loader(self, file_path: str) -> Optional[DocumentLoader]:
        """Get appropriate loader for file"""
        for loader in self.loaders:
            if loader.can_load(file_path):
                return loader
        
        return None
    
    def can_load(self, file_path: str) -> bool:
        """Check if any loader can handle this file"""
        return self.get_loader(file_path) is not None
    
    def load_document(self, file_path: str) -> Document:
        """Load a document using appropriate loader"""
        loader = self.get_loader(file_path)
        
        if not loader:
            raise ValueError(f"No loader available for: {file_path}")
        
        return loader.load(file_path)
    
    def get_supported_extensions(self) -> list:
        """Get list of supported file extensions"""
        return ['.txt', '.md', '.pdf', '.docx', '.html', '.htm']


# Global instance

_loader_registry = None

def get_loader_registry() -> LoaderRegistry:
    """Get or create global loader registry"""
    global _loader_registry
    if _loader_registry is None:
        _loader_registry = LoaderRegistry()
    return _loader_registry